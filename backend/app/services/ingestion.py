"""
Document ingestion service.
Handles parsing, chunking, embedding, and indexing documents into ChromaDB.
"""

import os
import logging
from pathlib import Path
from typing import List, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import chromadb
from sqlalchemy.orm import Session

from app.config import settings
from app.models import Document, DocumentChunk, DocumentStatus, PermissionLevel
from app.services.embeddings import embedding_service
from app.utils.chunking import chunk_markdown, chunk_plain_text

logger = logging.getLogger(__name__)


class IngestionService:
    """Handles document parsing, chunking, embedding, and indexing."""

    def __init__(self):
        self.chroma_client = chromadb.PersistentClient(path=settings.chroma_persist_dir)
        self.collection = self.chroma_client.get_or_create_collection(
            name=settings.chroma_collection_name,
            metadata={"hnsw:space": "cosine"},
        )

    def _parse_file(self, file_path: str, file_type: str) -> str:
        """Parse a document file and return its text content."""
        file_type = file_type.lower()

        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type == "docx":
            return self._parse_docx(file_path)
        elif file_type in ("txt", "md"):
            return self._parse_text(file_path)
        elif file_type == "html":
            return self._parse_html(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _parse_text(self, file_path: str) -> str:
        """Read plain text or markdown files."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_html(self, file_path: str) -> str:
        """Extract text from HTML using BeautifulSoup."""
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "lxml")
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text(separator="\n", strip=True)

    async def ingest_document(
        self,
        db: Session,
        file_path: str,
        title: str,
        department: str = "HR",
        owner: str = "HR Department",
        permission_level: str = "all_employees",
        version: str = "1.0",
    ) -> Document:
        """
        Full ingestion pipeline: parse → chunk → embed → index.
        """
        file_path = str(Path(file_path).resolve())
        filename = os.path.basename(file_path)
        file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

        # Create document record
        doc = Document(
            title=title,
            filename=filename,
            file_path=file_path,
            file_type=file_type,
            department=department,
            owner=owner,
            permission_level=PermissionLevel(permission_level),
            version=version,
            status=DocumentStatus.PARSING,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        try:
            # 1. Parse
            logger.info(f"Parsing document: {filename}")
            doc.status = DocumentStatus.PARSING
            db.commit()
            text = self._parse_file(file_path, file_type)

            if not text.strip():
                raise ValueError(f"No text extracted from {filename}")

            # 2. Chunk
            if file_type == "md" or self._looks_like_markdown(text):
                chunks = chunk_markdown(
                    text, filename,
                    max_chunk_size=settings.chunk_size,
                    overlap=settings.chunk_overlap,
                )
            else:
                chunks = chunk_plain_text(
                    text, filename,
                    max_chunk_size=settings.chunk_size,
                    overlap=settings.chunk_overlap,
                )

            # 3. Embed
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            doc.status = DocumentStatus.EMBEDDING
            db.commit()

            chunk_texts = [c["content"] for c in chunks]
            embeddings = await embedding_service.embed_batch(chunk_texts)

            # 4. Index into ChromaDB
            ids = []
            documents = []
            metadatas = []
            db_chunks = []

            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"{doc.id}_chunk_{i}"
                ids.append(chunk_id)
                documents.append(chunk["content"])
                metadatas.append({
                    "document_id": doc.id,
                    "document_title": title,
                    "section_title": chunk.get("section_title", ""),
                    "parent_section": chunk.get("parent_section", ""),
                    "department": department,
                    "permission_level": permission_level,
                    "filename": filename,
                    "chunk_index": str(i),
                })

                # Save chunk to SQLite too
                db_chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=doc.id,
                    chunk_index=str(i),
                    content=chunk["content"],
                    section_title=chunk.get("section_title", ""),
                    parent_section=chunk.get("parent_section", ""),
                    metadata_json=metadatas[-1],
                )
                db_chunks.append(db_chunk)

            # Batch upsert into ChromaDB
            self.collection.upsert(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas,
            )

            # Save chunks to SQLite
            db.add_all(db_chunks)
            doc.status = DocumentStatus.INDEXED
            doc.chunk_count = str(len(chunks))
            db.commit()

            logger.info(f"Successfully indexed '{title}' with {len(chunks)} chunks")
            return doc

        except Exception as e:
            logger.error(f"Ingestion failed for '{filename}': {e}")
            doc.status = DocumentStatus.FAILED
            doc.metadata_json = {"error": str(e)}
            db.commit()
            raise

    async def ingest_sample_data(self, db: Session) -> List[Document]:
        """Ingest all sample HR documents from the sample_data directory."""
        sample_dir = settings.sample_data_dir
        if not os.path.exists(sample_dir):
            raise FileNotFoundError(f"Sample data directory not found: {sample_dir}")

        documents = []
        sample_files = [
            ("employee_handbook.md", "Acme Corporation Employee Handbook"),
            ("leave_policy.md", "Acme Corporation Leave Policy"),
            ("onboarding_guide.md", "Acme Corporation New Employee Onboarding Guide"),
            ("benefits_overview.md", "Acme Corporation Employee Benefits Overview"),
            ("code_of_conduct.md", "Acme Corporation Code of Conduct"),
        ]

        for filename, title in sample_files:
            file_path = os.path.join(sample_dir, filename)
            if os.path.exists(file_path):
                # Skip if already ingested
                existing = db.query(Document).filter(Document.filename == filename).first()
                if existing:
                    logger.info(f"Skipping '{filename}' — already ingested")
                    documents.append(existing)
                    continue

                doc = await self.ingest_document(
                    db=db,
                    file_path=file_path,
                    title=title,
                    department="HR",
                    owner="HR Department",
                    permission_level="all_employees",
                )
                documents.append(doc)
            else:
                logger.warning(f"Sample file not found: {file_path}")

        return documents

    def _looks_like_markdown(self, text: str) -> bool:
        """Heuristic to detect if text is markdown-formatted."""
        lines = text.split("\n")[:20]
        md_indicators = sum(1 for line in lines if line.strip().startswith("#"))
        return md_indicators >= 2

    def get_collection_stats(self) -> dict:
        """Get ChromaDB collection statistics."""
        count = self.collection.count()
        return {
            "collection_name": settings.chroma_collection_name,
            "total_chunks": count,
        }

    def delete_document_chunks(self, document_id: str):
        """Remove all chunks for a document from ChromaDB."""
        try:
            results = self.collection.get(
                where={"document_id": document_id}
            )
            if results and results["ids"]:
                self.collection.delete(ids=results["ids"])
                logger.info(f"Deleted {len(results['ids'])} chunks for document {document_id}")
        except Exception as e:
            logger.error(f"Failed to delete chunks: {e}")


# Singleton instance
ingestion_service = IngestionService()
